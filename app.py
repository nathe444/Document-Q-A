import time
import streamlit as st
import os
from dotenv import load_dotenv
from langchain_groq import ChatGroq
from langchain.prompts import ChatPromptTemplate
from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
import tempfile
import docx

# Load environment variables from .env file
load_dotenv()
os.environ['GROQ_API_KEY'] = os.getenv('GROQ_API_KEY')

# Custom CSS styling for the Streamlit app
st.markdown(""" 
    <style>
        body {
            font-family: 'Roboto', sans-serif;
            background-color: #F7F9FC;  # Light background color for the app
        }

        h1 {
            color: #4A90E2;  # Primary color for the header
            font-size: 38px;
            font-weight: bold;
            margin-bottom: 20px;  # Margin below header
            text-align: center;  # Centered header
        }

        h2 {
            font-size: 26px;
            color: #333333;  # Darker color for subheaders
            margin-bottom: 10px;  # Margin below subheader
            text-align: center;  # Centered subheader
        }

        p {
            color: #5A6C7D;  # Color for paragraphs
            font-size: 18px;
            line-height: 1.6;  # Line height for readability
        }

        .stTextInput, .stFileUploader, .stButton {
            width: 100%;  # Full width input fields
            max-width: 700px;  # Max width for inputs
            margin: 0 auto;  # Centered inputs
        }
            
        .stFileUploader {
            margin-top: 10px;  # Space above file uploader
            padding: 10px;  # Padding for file uploader
        }

        .stButton {
            background-color: #4A90E2 !important;  # Button background color
            color: #ffffff !important;  # Button text color
            border-radius: 8px;  # Rounded corners for buttons
            padding: 12px !important;  # Button padding
            font-size: 16px;  # Button text size
            margin-top: 20px;  # Space above buttons
        }

        @media (max-width: 768px) {
            h1 {
                font-size: 28px;  # Responsive font size for header
            }

            h2 {
                font-size: 20px;  # Responsive font size for subheader
            }

            p {
                font-size: 16px;  # Responsive font size for paragraphs
            }

            .stTextInput, .stButton {
                padding: 10px;  # Responsive padding
                font-size: 14px;  # Responsive button text size
            }
        }

        .footer {
            text-align: center;  # Centered footer text
            color: #839ab2;  # Footer text color
            margin-top: 20px;  # Space above footer
            font-size: 16px;  # Footer text size
        }

        hr {
            margin-top: 110px;  # Space above horizontal rule
            border: none;  # No border for hr
            border-top: 1px solid #e0e0e0;  # Top border for hr
        }
    </style>
""", unsafe_allow_html=True)

# Main title and description for the app
st.markdown("<h1>📝 Document Q&A Chatbot</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align: center;'>Upload your document (PDF, DOCX, or TXT) and ask questions to get instant answers!</p>", unsafe_allow_html=True)

# File uploader for documents
uploaded_file = st.file_uploader("Upload your document (PDF, DOCX, or TXT)", type=['pdf', 'docx', 'txt'], label_visibility='collapsed')

st.markdown("<div style='margin-bottom: 20px;'></div>", unsafe_allow_html=True)

# Document class to hold document content and metadata
class Document:
    def __init__(self, content):
        self.page_content = content  # Document content
        self.metadata = {}  # Metadata dictionary

# Function to extract text from DOCX files
def extract_text_from_docx(docx_file):
    doc = docx.Document(docx_file)  # Load DOCX file
    extracted_text = '\n'.join([para.text for para in doc.paragraphs])  # Extract text from paragraphs
    return extracted_text  # Return extracted text

documents = []  # List to hold documents
if uploaded_file:
    with st.spinner('Processing document...'):  # Show loading spinner while processing
        # Handling PDF file upload
        if uploaded_file.type == 'application/pdf':
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                temp_file.write(uploaded_file.read())  # Write uploaded file to temp file
                temp_file_path = temp_file.name  # Get temp file path
            loader = PyPDFLoader(temp_file_path)  # Initialize PDF loader
            documents = loader.load()  # Load documents from PDF
        # Handling DOCX file upload
        elif uploaded_file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':  # DOCX type
            extracted_text = extract_text_from_docx(uploaded_file)  # Extract text from DOCX
            documents = [Document(extracted_text)]  # Create Document instance
        # Handling plain text file upload
        elif uploaded_file.type == 'text/plain':
            text = uploaded_file.read().decode("utf-8")  # Decode uploaded text file
            documents = [Document(text)]  # Create Document instance

    # Proceed if documents are processed
    if documents:
        with st.spinner('Splitting documents and embedding...'):  # Show loading spinner while processing
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)  # Initialize text splitter
            final_documents = text_splitter.split_documents(documents)  # Split documents into chunks

            embeddings = HuggingFaceEmbeddings()  # Initialize embeddings
            vectors = FAISS.from_documents(final_documents, embeddings)  # Create vector store from documents and embeddings

            llm = ChatGroq(temperature=0.5, model_name="mixtral-8x7b-32768")  # Initialize language model

            prompt = ChatPromptTemplate.from_template(  # Define prompt template
                """
                Answer the question based on the provided context.
                <context>
                {context}
                </context>
                Question: {input}
                """
            )

            document_chain = create_stuff_documents_chain(llm, prompt)  # Create document chain for processing
            retriever = vectors.as_retriever()  # Create retriever from vector store
            retrieval_chain = create_retrieval_chain(retriever, document_chain)  # Create retrieval chain

        st.success("Document processed! Now ask your questions.")  # Notify user that document is processed

# Text input for user to ask questions
question = st.text_input("Ask a question about your document:")

if question:  # If a question is asked
    if documents:  # If documents are available
        start = time.process_time()  # Start timer for response generation
        with st.spinner('Generating answer...'):  # Show loading spinner while generating answer
            response = retrieval_chain.invoke({'input': question})  # Invoke retrieval chain with question
        st.write(f"**Answer:** {response['answer']}")  # Display the answer
        st.write(f"_Response generated in {time.process_time() - start:.2f} seconds_")  # Display response time

        with st.expander("View related document context"):  # Expandable section to view context
            for i, doc in enumerate(response['context']):  # Iterate through context documents
                st.write(f"**Context {i+1}:** {doc.page_content}")  # Display each context
                st.write("---")  # Separator for contexts
    else:
        st.write("**Note:** Please upload a document to ask document-related questions.")  # Prompt user to upload document if none is available

# Footer section for the app
st.markdown("<hr><p class='footer'>Built with 💙 using LangChain & Groq</p>", unsafe_allow_html=True)
